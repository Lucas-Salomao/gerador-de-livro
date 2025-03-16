import os
from typing import Dict, List, Tuple, Any, TypedDict, Optional
import json
from pathlib import Path
import logging
from dotenv import load_dotenv

# Bibliotecas para LangGraph
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
import langchain

# Bibliotecas para Gemini/Vertex AI
from google.cloud import aiplatform
from vertexai.generative_models import GenerativeModel, Part
import vertexai

import google.generativeai as genai

# Biblioteca para exportação
from docx import Document

load_dotenv()

# Configuração de logs
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("book_generation.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Inicializar Gemini API
def init_gemini_api(api_key: str):
    """Inicializa a conexão com o Gemini API."""
    logger.info("Inicializando Gemini API...")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-2.0-flash')

# Função auxiliar para parsing seguro de JSON
def safe_json_parse(response_text: str, fallback: Any) -> Any:
    """Tenta decodificar JSON e retorna um fallback em caso de erro."""
    if response_text.startswith("```json"):
        response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
    elif response_text.startswith("```"):
        response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
    
    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        logger.error(f"Erro ao decodificar JSON: {response_text[:100]}... Usando fallback.")
        return fallback

# Definição dos estados do grafo
class BookState(TypedDict, total=False):
    theme: str
    title: str
    genre: str
    target_audience: str
    outline: List[Dict[str, Any]]
    chapters: Dict[int, Dict[str, str]]
    current_chapter: int
    status: str
    feedback: str
    export_path: str
    feedback_path: str

# Funções para cada etapa do processo
def get_book_info(state: BookState, model) -> Dict[str, Any]:
    """Obtém informações básicas e gera o título com base no tema."""
    logger.info(f"Estado recebido em get_book_info: {state}")
    logger.info("Coletando informações básicas do livro...")
    updates = {}
    
    theme = state.get("theme", "Um tema genérico")
    genre = state.get("genre", "Ficção")
    target_audience = state.get("target_audience", "Adultos")
    
    updates["theme"] = theme
    updates["genre"] = genre
    updates["target_audience"] = target_audience
    
    prompt = f"""
    Você é um especialista em redação técnica. Baseado no seguinte tema, gênero e público-alvo, sugira um título formal e técnico que reflita um enfoque analítico e informativo:
    Tema: {theme}
    Gênero: {genre}
    Público-Alvo: {target_audience}
    Responda SOMENTE em formato JSON com a chave "title", sem texto adicional. Exemplo: {{"title": "Fundamentos de Exploração Espacial"}}. Não inclua bloco de código, ou seja ```json```
    O Título deve ter no máximo 80 caracteres. Caracteres inválidos para o título: \ / : * ? " < > |
    """
    
    logger.info("Gerando título com base no tema...")
    response = model.generate_content(prompt)
    logger.debug(f"Resposta bruta do modelo: {response.text}")
    info = safe_json_parse(response.text, {"title": f"Livro sobre {theme}"})
    updates["title"] = info.get("title", f"Livro sobre {theme}")
    logger.info(f"Título gerado: {updates['title']}")
    
    updates["status"] = "book_info_collected"
    logger.debug(f"Atualizações de get_book_info: {updates}")
    return updates

def create_outline(state: BookState, model) -> Dict[str, Any]:
    """Cria o sumário do livro baseado nas informações fornecidas."""
    logger.info("Criando sumário do livro...")
    
    prompt = f"""
    Você é um especialista técnico elaborando um livro técnico para estudo de um determinado tema. 
    Baseado nas seguintes informações, crie um sumário detalhado e com pelo menos 3 níveis de aprofundamento com foco em aspectos técnicos e práticos:
    
    Tema: {state['theme']}
    Título sugerido: {state['title']}
    Gênero: {state['genre']}
    Público-Alvo: {state['target_audience']}
    
    Inclua entre 5 a 20 capítulos, cada um abordando um aspecto técnico ou prático do tema, com títulos objetivos e descrições que detalhem o conteúdo analítico a ser explorado.
    Responda SOMENTE em formato JSON com uma lista de objetos contendo "chapter_number", "chapter_title" e "chapter_description".
    Exemplo: [{{"chapter_number": 1, "chapter_title": "Princípios de Propulsão Espacial", "chapter_description": "Análise dos sistemas de propulsão usados em missões espaciais"}}]
    Não inclua bloco de código, ou seja ```json```
    """
    
    response = model.generate_content(prompt)
    logger.debug(f"Resposta bruta do modelo: {response.text}")
    outline_data = safe_json_parse(response.text, [
        {"chapter_number": 1, "chapter_title": "Introdução", 
         "chapter_description": f"Exploração inicial do tema {state['theme']}."}
    ])
    
    if len(outline_data) < 5:
        logger.warning("Sumário com menos de 5 capítulos. Adicionando capítulos extras.")
        for i in range(len(outline_data) + 1, 6):
            outline_data.append({
                "chapter_number": i,
                "chapter_title": f"Capítulo {i}",
                "chapter_description": f"Continuação da exploração de {state['theme']}."
            })
    
    updates = {
        "outline": outline_data,
        "chapters": {item["chapter_number"]: {"title": item["chapter_title"], 
                                              "description": item["chapter_description"],
                                              "content": ""} 
                     for item in outline_data},
        "current_chapter": 1,
        "status": "outline_created"
    }
    logger.info(f"Sumário criado com {len(outline_data)} capítulos.")
    logger.info(f"Capítulos gerados: {updates['chapters']}")
    return updates

def write_chapter(state: BookState, model, st_session=None) -> Dict[str, Any]:
    """Escreve o conteúdo para o capítulo atual."""
    current = state["current_chapter"]
    updates = {}
    if current > len(state["chapters"]):
        updates["status"] = "all_chapters_written"
        logger.info("Todos os capítulos foram escritos.")
        return updates
    
    chapter_info = state["chapters"][current]
    logger.info(f"Escrevendo Capítulo {current}: {chapter_info['title']}...")
    
    # Notificar o Streamlit sobre o progresso
    if st_session:
        st_session.write(f"Gerando Capítulo {current}: {chapter_info['title']}")
        progress = int((current / len(state["chapters"])) * 100)
        st_session.progress(progress)
    
    prev_content = ""
    if current > 1 and state["chapters"].get(current-1, {}).get("content"):
        prev_chapter = state["chapters"][current-1]
        prev_content = f"""
        Resumo do capítulo anterior ({current-1}: {prev_chapter['title']}):
        {prev_chapter['content'][:500]}... (resumido)
        """
    
    prompt = f"""
    Você é um especialista técnico escrevendo um livro intitulado "{state['title']}" 
    com o tema "{state['theme']}" no gênero "{state['genre']}" para o público "{state['target_audience']}".
    
    Escreva o Capítulo {current}: "{chapter_info['title']}".
    
    Descrição do capítulo: {chapter_info['description']}
    
    {prev_content}
    
    Escreva um texto técnico e analítico, com linguagem formal e objetiva. Inclua informações técnicas detalhadas, exemplos contextualizados (reais ou hipotéticos), dados relevantes e explicações claras. Evite diálogos narrativos ou descrições literárias excessivas. Estruture o conteúdo com seções claras (ex.: introdução, análise, exemplos, conclusão). O capítulo deve ter pelo menos 3000 palavras. Seja o mais detalhista possível e aborde o tema do capítulo com profundidade e bastante exemplo.
    """
    
    response = model.generate_content(prompt)
    updated_chapters = state["chapters"].copy()
    updated_chapters[current]["content"] = response.text
    updates["chapters"] = updated_chapters
    logger.info(f"Capítulo {current} concluído com sucesso.")
    
    # Exibir o conteúdo gerado no Streamlit
    if st_session:
        # st_session.write(f"### Capítulo {current}: {chapter_info['title']}")
        st_session.write(response.text)
    
    updates["current_chapter"] = current + 1
    updates["status"] = "chapter_written" if updates["current_chapter"] <= len(state["chapters"]) else "all_chapters_written"
    return updates

def review_and_edit(state: BookState, model) -> Dict[str, Any]:
    """Revisa e edita o livro completo."""
    logger.info("Revisando e editando o livro...")
    book_summary = f"""
    Tema: {state['theme']}
    Título: {state['title']}
    Gênero: {state['genre']}
    Público-alvo: {state['target_audience']}
    
    Sumário:
    """
    for chapter_num, chapter_data in sorted(state["chapters"].items()):
        book_summary += f"\nCapítulo {chapter_num}: {chapter_data['title']} - {chapter_data['description'][:100]}..."
    
    prompt = f"""
    Você é um editor revisando o livro:
    
    {book_summary}
    
    Forneça feedback sobre estrutura, fluxo narrativo, consistência com o tema "{state['theme']}" 
    e apelo ao público-alvo. Sugira melhorias. Revise tecnicamente o livro e verifique se há alguma inconsistência.
    """
    
    response = model.generate_content(prompt)
    updates = {
        "feedback": response.text,
        "status": "reviewed"
    }
    logger.info("Revisão concluída. Feedback gerado.")
    
    return updates

def export_feedback(state: BookState) -> Dict[str, Any]:
    """Exporta o feedback para um arquivo TXT."""
    logger.info("Exportando feedback para TXT...")
    feedback_path = f"{state['title'].replace(' ', '_')}_feedback.txt"
    with open(feedback_path, "w", encoding="utf-8") as f:
        f.write(state["feedback"])
    updates = {
        "feedback_path": feedback_path,
        "status": "feedback_exported"
    }
    logger.info(f"Feedback exportado com sucesso para: {feedback_path}")
    return updates

def export_book(state: BookState) -> Dict[str, Any]:
    """Exporta o livro apenas para DOCX."""
    logger.info("Exportando livro para DOCX...")
    doc = Document()
    doc.add_heading(state["title"], 0)
    doc.add_paragraph(f"Tema: {state['theme']}")
    doc.add_paragraph(f"Gênero: {state['genre']}")
    doc.add_paragraph(f"Público-alvo: {state['target_audience']}")
    
    for chapter_num, chapter_data in sorted(state["chapters"].items()):
        doc.add_heading(f"Capítulo {chapter_num}: {chapter_data['title']}", 1)
        doc.add_paragraph(chapter_data["content"])
    
    doc_path = f"{state['title'].replace(' ', '_')}.docx"
    doc.save(doc_path)
    updates = {
        "export_path": doc_path,
        "status": "exported"
    }
    logger.info(f"Livro exportado com sucesso para: {doc_path}")
    return updates

def router(state: BookState) -> str:
    """Decide o próximo estado."""
    status_map = {
        "start": "get_book_info",
        "book_info_collected": "create_outline",
        "outline_created": "write_chapter",
        "chapter_written": "write_chapter",
        "all_chapters_written": "review_and_edit",
        "reviewed": "export_feedback",
        "feedback_exported": "export_book",
        "exported": END
    }
    next_state = status_map.get(state["status"], END)
    logger.debug(f"Transição de estado: {state['status']} -> {next_state}")
    return next_state

def create_book_agent(model, st_session=None):
    """Cria o agente de geração de livros."""
    logger.info("Criando agente de geração de livros...")
    workflow = StateGraph(BookState)
    
    workflow.add_node("get_book_info", lambda state: get_book_info(state, model))
    workflow.add_node("create_outline", lambda state: create_outline(state, model))
    workflow.add_node("write_chapter", lambda state: write_chapter(state, model, st_session))
    workflow.add_node("review_and_edit", lambda state: review_and_edit(state, model))
    workflow.add_node("export_feedback", export_feedback)
    workflow.add_node("export_book", export_book)
    
    workflow.set_entry_point("get_book_info")
    
    workflow.add_conditional_edges("get_book_info", router)
    workflow.add_conditional_edges("create_outline", router)
    workflow.add_conditional_edges("write_chapter", router)
    workflow.add_conditional_edges("review_and_edit", router)
    workflow.add_conditional_edges("export_feedback", router)
    workflow.add_conditional_edges("export_book", router)
    
    memory = MemorySaver()
    return workflow.compile(checkpointer=memory)

def main(custom_theme: str = "", custom_genre: str = "", custom_audience: str = "", st_session=None):
    """Executa o agente de geração de livros."""
    logger.info("Iniciando processo de geração de livro...")
    try:
        model = init_gemini_api(os.getenv("GEMINI_API_KEY"))
        book_agent = create_book_agent(model, st_session)
        
        initial_state = BookState(status="start")
        if custom_theme:
            initial_state["theme"] = custom_theme
        if custom_genre:
            initial_state["genre"] = custom_genre
        if custom_audience:
            initial_state["target_audience"] = custom_audience
        
        config = {"configurable": {"thread_id": "1"}}
        
        for output in book_agent.stream(initial_state, config=config):
            node_name = list(output.keys())[0] if output else "unknown"
            stage = output.get(node_name, {}).get("status", "desconhecido")
            if st_session:
                st_session.write(f"Concluído: {stage}")
            if stage == "book_info_collected":
                if st_session:
                    st_session.write(f"**Tema:** {output[node_name]['theme']}")
                    st_session.write(f"**Título gerado:** {output[node_name]['title']}")
                    st_session.write(f"**Gênero:** {output[node_name]['genre']}")
                    st_session.write(f"**Público-alvo:** {output[node_name]['target_audience']}")
            elif stage == "outline_created":
                if st_session:
                    st_session.write(f"Sumário criado com {len(output[node_name]['outline'])} capítulos")
            elif stage == "chapter_written":
                if st_session:
                    st_session.write(f"Capítulo {output[node_name]['current_chapter']-1} concluído")
            elif stage == "feedback_exported":
                if st_session:
                    st_session.write(f"Feedback exportado para: {output[node_name]['feedback_path']}")
            elif stage == "exported":
                if st_session:
                    st_session.write(f"Livro exportado para: {output[node_name]['export_path']}")
        
        final_state = book_agent.checkpointer.get(config)
        logger.info("Processo de geração de livro concluído!")
        return final_state
    except Exception as e:
        logger.error(f"Erro durante a geração do livro: {e}")
        if st_session:
            st_session.error(f"Erro durante a geração do livro: {e}")
        return {"status": "error", "message": str(e)}